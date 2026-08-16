import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import datetime

st.set_page_config(
    page_title="Smart Document Writing Tool",
    page_icon="📋",
    layout="centered",
    initial_sidebar_state="expanded"
)

# ── Custom CSS with "For Aiswarya" Golden Tag in Header ───────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+Malayalam:wght@400;600;700&family=Noto+Sans+Malayalam:wght@400;500;600&display=swap');

html, body, [data-testid="stAppViewContainer"] {
    background: #f0ebe0 !important;
    font-family: 'Noto Sans Malayalam', sans-serif;
    color: #2C3E50 !important;
}

#MainMenu, footer, header { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }

.block-container {
    max-width: 800px !important;
    padding: 2rem 1.5rem 4rem !important;
}

/* App Header */
.app-header {
    background: linear-gradient(135deg, #1A4D2E 0%, #0F2C1A 100%);
    border-radius: 10px;
    padding: 22px 28px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    position: relative;
    box-shadow: 0 4px 20px rgba(26,77,46,.25);
}
.app-header-left {
    display: flex;
    align-items: center;
    gap: 16px;
}
.app-header-icon {
    font-size: 2rem;
    background: rgba(255,255,255,.15);
    width: 52px; height: 52px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    flex-shrink: 0;
}
.app-header h1 {
    margin: 0; color: white;
    font-family: 'Noto Serif Malayalam', serif;
    font-size: 1.3rem; font-weight: 700; line-height: 1.4;
}
.app-header p { 
    margin: 4px 0 0; 
    color: rgba(255,255,255,.85); 
    font-size: .82rem; 
}
.app-header-tag {
    position: absolute;
    bottom: 8px;
    right: 18px;
    color: #FFD700;
    font-size: 0.75rem;
    font-style: italic;
    font-weight: 600;
    letter-spacing: 0.5px;
}

/* Card Design */
.form-card {
    background: white;
    border-radius: 10px;
    padding: 24px 28px;
    margin-bottom: 20px;
    box-shadow: 0 2px 12px rgba(0,0,0,.07);
    border-top: 4px solid #1A4D2E;
}
.form-section-title {
    font-family: 'Noto Serif Malayalam', serif;
    font-size: 1.05rem;
    font-weight: 700;
    color: #1A4D2E;
    margin-bottom: 14px;
    padding-bottom: 8px;
    border-bottom: 2px solid #e2ebd8;
}

/* Streamlit Labels High Contrast */
.stSelectbox label, .stTextArea label, .stTextInput label {
    color: #1A4D2E !important;
    font-weight: 600 !important;
}

/* Input Fields Styling - Fixed Text Color to Dark */
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
[data-testid="stTextInput"] input {
    border-radius: 6px !important;
    border-color: #b8c5b9 !important;
    font-family: 'Noto Sans Malayalam', sans-serif !important;
    font-size: 0.95rem !important;
    color: #111111 !important;
    background-color: #ffffff !important;
}

/* Dropdown selected text color fix */
[data-testid="stSelectbox"] span {
    color: #111111 !important;
}

[data-testid="stTextArea"] textarea:focus,
[data-testid="stTextInput"] input:focus {
    border-color: #1A4D2E !important;
    box-shadow: 0 0 0 2px rgba(26,77,46,0.15) !important;
}

/* Primary Button */
[data-testid="stButton"] > button[kind="primary"] {
    background: linear-gradient(135deg, #1A4D2E, #2A7347) !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Noto Sans Malayalam', sans-serif !important;
    font-size: 1.05rem !important;
    font-weight: 600 !important;
    height: 52px !important;
    box-shadow: 0 4px 14px rgba(26,77,46,.3) !important;
    width: 100% !important;
    color: white !important;
}

/* Sidebar styling */
[data-testid="stSidebar"] {
    background-color: #11301c;
}
[data-testid="stSidebar"] * {
    color: #ffffff !important;
}
.sidebar-info {
    font-size: 0.85rem;
    line-height: 1.6;
    color: #d1e8d5 !important;
    background: rgba(0,0,0,0.25);
    padding: 15px;
    border-radius: 8px;
    margin-top: 15px;
    border-left: 3px solid #81c784;
}
.sidebar-info a {
    color: #81c784 !important;
}
</style>
""", unsafe_allow_html=True)

# ── Document Types Dictionary ────────────────────────────────────────────
DOC_GROUPS = {
    "📝 അപേക്ഷകൾ (Applications - Public to Office)": {
        "app_general": "പൊതു അപേക്ഷ (General Application)",
        "app_income": "വരുമാന Certificate അപേക്ഷ",
        "app_nativity": "ജനന/നാട്ടുകാർ Certificate",
        "app_residence": "താമസ Certificate",
        "app_caste": "ജാതി Certificate",
        "app_building": "കെട്ടിട അനുമതി അപേക്ഷ",
        "app_pension": "പെൻഷൻ/ആനുകൂല്യം",
        "app_leave": "അവധി അപേക്ഷ (Leave Letter)",
        "app_complaint": "പരാതി (Complaint)"
    },
    "📨 ഔദ്യോഗിക കത്തുകൾ (Official Letters - Office/Public)": {
        "letter": "ഔദ്യോഗിക കത്ത് (Office to Office / Public)",
        "do_letter": "അർദ്ധ-ഔദ്യോഗിക കത്ത് (D.O. Letter)",
        "forwarding": "അയക്കൽ കത്ത് (Forwarding Letter)",
        "reminder": "ഓർമ്മപ്പെടുത്തൽ (Reminder Letter)"
    },
    "📜 ഉത്തരവുകൾ & അറിയിപ്പുകൾ (Orders & RTI)": {
        "order": "ഉത്തരവ് (Government Order)",
        "circular": "സർക്കുലർ (Circular)",
        "public_notice": "പൊതു അറിയിപ്പ് (Public Notice)",
        "rti_reply": "വിവരാവകാശ രേഖ / അപേക്ഷ (RTI)"
    }
}

ALL_TYPES = {k: v for g in DOC_GROUPS.values() for k, v in g.items()}

# ── Session State Initialization ─────────────────────────────────────────
for k, v in {
    'output': '', 'edit_mode': False,
    'docx_cache': None,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Sidebar: API Key & Info ──────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🔑 API Key Setup")
    secret_key = st.secrets.get("GEMINI_API_KEY", "")
    
    user_api_key = st.text_input(
        "Google Gemini API Key", 
        type="password", 
        value=secret_key,
        placeholder="AIzaSy..."
    )
    active_api_key = user_api_key or secret_key

    if active_api_key:
        st.success("✅ API Key Ready!")
    else:
        st.warning("⚠️ API Key നൽകുക.")

    st.markdown("""
    <div class="sidebar-info">
        <h3 style="margin-bottom: 8px; color: #ffffff !important; font-size: 0.95rem;">❓ API Key എങ്ങനെ ലഭിക്കും?</h3>
        <ol style="padding-left: 15px; margin-bottom: 0;">
            <li><a href="https://aistudio.google.com/app/apikey" target="_blank">Google AI Studio</a> സന്ദർശിക്കുക.</li>
            <li>Google Account വഴി ലോഗിൻ ചെയ്യുക.</li>
            <li>"Create API Key" ക്ലിക്ക് ചെയ്ത് കോപ്പി ചെയ്യുക.</li>
        </ol>
    </div>
    """, unsafe_allow_html=True)

# ── Master Prompt Builder ────────────────────────────────────────────────
def build_master_prompt(group_sel, doc_type_key, sender_info, recipient_info, details, language):
    doc_label = ALL_TYPES[doc_type_key]
    
    if language == "English":
        lang_instruction = "professional and official English"
        system_role = "You are an expert AI Assistant specialized in drafting official government documents, applications, and letters."
    else:
        lang_instruction = "ഔദ്യോഗിക ഭരണമലയാള ശൈലിയിൽ (Official Malayalam)"
        system_role = "നീ കേരള സർക്കാരിന്റെ ഔദ്യോഗിക ഫയലുകൾ, അപേക്ഷകൾ, കത്തുകൾ എന്നിവ തയ്യാറാക്കുന്നതിൽ അഗാധമായ അറിവുള്ള ഒരു 'Senior Section Officer' ആണ്."

    return f"""{system_role}
തിരഞ്ഞെടുത്ത വിഭാഗം: {group_sel}
രേഖയുടെ തരം: '{doc_label}'

നൽകിയിട്ടുള്ള വിവരങ്ങൾ പരിശോധിക്കുക:
1. കത്ത് അയക്കുന്ന ആളുടെ/സ്ഥാപനത്തിന്റെ വിവരങ്ങൾ (Sender Info): {sender_info}
2. കത്ത് ലഭിക്കേണ്ട ആളുടെ/സ്ഥാപനത്തിന്റെ വിവരങ്ങൾ (Recipient Info): {recipient_info}
3. വിഷയവും വിവരങ്ങളും (Subject & Details): {details}

പ്രധാന നിർദ്ദേശങ്ങൾ:
- കത്ത/അപേക്ഷ ഏതുതരം (Public to Office, Office to Office, Office to Public) ആണെന്ന് ഈ വിവരങ്ങളിൽ നിന്ന് മനസ്സിലാക്കി അതിനനുയോജ്യമായ ഔദ്യോഗിക ഫോർമാറ്റ് സ്വീകരിക്കുക.
- ഭാഷ: പൂർണ്ണമായും {lang_instruction} ആയിരിക്കണം.
- ഔട്ട്പുട്ടിൽ തയ്യാറാക്കിയ പൂർണ്ണമായ രേഖ മാത്രമേ ഉണ്ടാകാവൂ. യാതൊരുവിധ Markdown ചിഹ്നങ്ങളോ (** ##) അനാവശ്യ വിവരണങ്ങളും പാടില്ല.
"""

def call_gemini(api_key, model_name, prompt):
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=model_name,
        contents=prompt,
        config=types.GenerateContentConfig(temperature=0.3)
    )
    return response.text.strip()

def make_docx(text, label):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin = Cm(2.5)
    
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(f"[ {label} ]")
    tr.font.size = Pt(8); tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1A, 0x4D, 0x2E)
    tr.font.name = "Noto Serif Malayalam"
    pPr = tp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'4');   bot.set(qn('w:color'),'1A4D2E')
    pBdr.append(bot); pPr.append(pBdr)
    
    doc.add_paragraph()
    
    for line in text.split('\n'):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11); r.font.name = "Noto Serif Malayalam"
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:line'),'360'); sp.set(qn('w:lineRule'),'auto')
        p._p.get_or_add_pPr().append(sp)
    
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ── Main UI Header with "For Aiswarya" Tag ──────────────────────────────
st.markdown("""
<div class="app-header">
  <div class="app-header-left">
    <div class="app-header-icon">📋</div>
    <div>
      <h1>Smart Document Writing Tool</h1>
      <p>അപേക്ഷ, കത്തുകൾ, ഉത്തരവുകൾ &nbsp;·&nbsp; AI Powered</p>
    </div>
  </div>
  <div class="app-header-tag">For Aiswarya</div>
</div>
""", unsafe_allow_html=True)

# ── Form Input Section ───────────────────────────────────────────────────
st.markdown('<div class="form-card">', unsafe_allow_html=True)
st.markdown('<div class="form-section-title">📋 രേഖയുടെ തരം തിരഞ്ഞെടുക്കുക</div>', unsafe_allow_html=True)

col_g, col_t, col_l = st.columns([1.2, 1.2, 0.9])
with col_g:
    group_sel = st.selectbox("വിഭാഗം (Category)", list(DOC_GROUPS.keys()))
with col_t:
    type_map = DOC_GROUPS[group_sel]
    doc_type = st.selectbox("രേഖയുടെ തരം (Type)", list(type_map.keys()), format_func=lambda x: type_map[x])
with col_l:
    doc_language = st.selectbox("ഭാഷ (Language)", ["Malayalam", "English"])

st.markdown('<div class="form-section-title" style="margin-top: 25px;">📝 വിലാസവും വിവരങ്ങളും നൽകുക</div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)
with c1:
    sender_info = st.text_area(
        "കത്ത് അയക്കുന്ന ആളുടെ/സ്ഥാപനത്തിന്റെ വിവരങ്ങൾ (Sender Info)", 
        placeholder="ഉദാ: പൊതുജനം അല്ലെങ്കിൽ ഒരു ഓഫീസിന്റെ പേര്\nപേര് / ഓഫീസിന്റെ പേര്\nവിലാസം\nഫോൺ നമ്പർ", 
        height=120
    )
with c2:
    recipient_info = st.text_area(
        "കത്ത് ലഭിക്കേണ്ട ആളുടെ/സ്ഥാപനത്തിന്റെ വിവരങ്ങൾ (Recipient Info)", 
        placeholder="ഉദാ: പഞ്ചായത്ത് സെക്രട്ടറി / മറ്റൊരു ഓഫീസർ\nപദവി\nഓഫീസിന്റെ പേര്\nസ്ഥലം", 
        height=120
    )

details = st.text_area(
    "വിഷയവും മറ്റ് വിവരങ്ങളും (Subject & Details)",
    placeholder="ഉദാ: കുടിവെള്ള കണക്ഷൻ ലഭിക്കുന്നത് സംബന്ധിച്ച്... അല്ലെങ്കിൽ ഒരു ഓഫീസിൽ നിന്ന് മറ്റേ ഓഫീസിലേക്ക് അയക്കേണ്ട കത്തിന്റെ വിവരങ്ങൾ...",
    height=150
)

model_name = st.text_input(
    "AI Model Name (ഫ്യൂച്ചർ അപ്ഡേറ്റുകൾക്കായി ടൈപ്പ് ചെയ്യാം)",
    value="gemini-2.5-flash",
    help="ഉദാഹരണത്തിന് gemini-2.5-flash അല്ലെങ്കിൽ gemini-2.5-pro എന്ന് നൽകാം."
)

st.markdown('</div>', unsafe_allow_html=True)

# ── Generate Button ───────────────────────────────────────────────────────
if st.button("⚡ മികച്ച ഔദ്യോഗിക രേഖ തയ്യാറാക്കുക", type="primary"):
    if not active_api_key:
        st.error("⚠️ ദയവായി ഇടത് വശത്തെ സൈഡ്‌ബാറിൽ Google Gemini API Key നൽകുക.")
    elif not details.strip():
        st.error("⚠️ ദയവായി വിഷയവും വിവരങ്ങളും നൽകുക.")
    else:
        with st.spinner("AI ഭരണഭാഷയിൽ രേഖ തയ്യാറാക്കുന്നു..."):
            try:
                prompt = build_master_prompt(group_sel, doc_type, sender_info, recipient_info, details, doc_language)
                result = call_gemini(active_api_key, model_name, prompt)
                
                final_label = type_map[doc_type] if doc_language == "Malayalam" else "Official Document"

                st.session_state.output = result
                st.session_state.edit_mode = False
                st.session_state.doc_label = final_label
                st.session_state.docx_cache = make_docx(result, final_label)
                st.rerun()
            except Exception as e:
                error_msg = str(e)
                if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                    st.error("⚠️ API Quota Limit കഴിഞ്ഞിരിക്കുന്നു. ദയവായി കുറച്ചുസമയം കഴിഞ്ഞ് വീണ്ടും ശ്രമിക്കുക.")
                else:
                    st.error(f"❌ Error: {error_msg}")

# ── Output Section ────────────────────────────────────────────────────────
if st.session_state.output:
    doc_label = st.session_state.doc_label or "Official Document"

    if st.session_state.edit_mode:
        _docx_bytes = make_docx(st.session_state.output, doc_label)
    else:
        _docx_bytes = st.session_state.docx_cache or make_docx(st.session_state.output, doc_label)

    st.markdown('<div class="form-card">', unsafe_allow_html=True)
    st.markdown('<div class="form-section-title">📄 തയ്യാറാക്കിയ രേഖ (Output Preview)</div>', unsafe_allow_html=True)

    col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
    if col_btn1.button("✏️ Edit" if not st.session_state.edit_mode else "👁️ Preview", use_container_width=True):
        st.session_state.edit_mode = not st.session_state.edit_mode
        st.rerun()

    col_btn2.download_button(
        "⬇️ .docx",
        data=_docx_bytes,
        file_name=f"document-{datetime.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    if col_btn3.button("🖨️ Print", use_container_width=True):
        st.session_state['do_print'] = True
    if col_btn4.button("🗑️ Clear", use_container_width=True):
        st.session_state.output = ''
        st.session_state.edit_mode = False
        st.rerun()

    if st.session_state.edit_mode:
        st.caption("✏️ താഴെ നേരിട്ട് തിരുത്തലുകൾ (Edit) വരുത്താം:")
        edited = st.text_area(
            "edit",
            value=st.session_state.output,
            height=450,
            label_visibility="collapsed"
        )
        st.session_state.output = edited
    else:
        print_js = ""
        if st.session_state.get('do_print'):
            print_js = "<script>window.print();</script>"
            st.session_state['do_print'] = False

        safe_text = (st.session_state.output
                     .replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))

        st.components.v1.html(f"""
<!DOCTYPE html><html><head><meta charset="UTF-8">
<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+Malayalam:wght@400;600;700&display=swap" rel="stylesheet">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:#f0ebe0;padding:8px;font-size:14px}}
.sheet{{
  background:white;border-radius:6px;
  padding:40px 48px;
  font-family:'Noto Serif Malayalam',serif;
  font-size:14px;line-height:2.2;color:#111;
  white-space:pre-wrap;
  box-shadow:0 4px 15px rgba(0,0,0,.08);
  position:relative;
  border-top: 4px solid #1A4D2E;
}}
@media print{{
  body{{background:white;padding:0}}
  .sheet{{box-shadow:none;border:none;padding:0}}
}}
</style></head>
<body><div class="sheet">{safe_text}</div>{print_js}</body></html>
""", height=550, scrolling=True)

        word_count = len(st.session_state.output.split())
        st.caption(f"📊 {word_count} words · {len(st.session_state.output)} characters")
    
    st.markdown('</div>', unsafe_allow_html=True)
